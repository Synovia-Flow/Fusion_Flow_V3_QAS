"""Generate Fusion Flow SOP as a Word document with embedded screenshots."""
import os, sys, subprocess

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

ROOT    = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SS_DIR  = os.path.join(ROOT, "docs", "sop", "screenshots")
OUT     = os.environ.get("SOP_OUT") or os.path.join(ROOT, "docs", "sop", "Fusion_Flow_V2_BKD_SOP.docx")

DARK_BLUE  = RGBColor(0x0D, 0x22, 0x47)
MID_BLUE   = RGBColor(0x1A, 0x40, 0x70)
LINK_BLUE  = RGBColor(0x1A, 0x6B, 0xBA)
GREY       = RGBColor(0x5A, 0x6A, 0x84)
GREEN      = RGBColor(0x06, 0x5F, 0x46)
RED        = RGBColor(0x99, 0x1B, 0x1B)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = DARK_BLUE
    p.runs[0].font.size = Pt(22)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = MID_BLUE
    p.runs[0].font.size = Pt(16)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = MID_BLUE
    p.runs[0].font.size = Pt(13)
    return p

def body(text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def note(text, danger=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(("⚠ " if danger else "ℹ ") + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RED if danger else RGBColor(0x92, 0x40, 0x0E)
    run.font.italic = True
    return p

def screenshot(filename, caption):
    path = os.path.join(SS_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6.2))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = GREY
            run.font.italic = True
    else:
        doc.add_paragraph(f"[Screenshot not found: {filename}]")

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = DARK_BLUE
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
    return table

def step(num, title, detail):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"  {num}.  ")
    r1.bold = True
    r1.font.color.rgb = LINK_BLUE
    r1.font.size = Pt(11)
    r2 = p.add_run(f"{title} — ")
    r2.bold = True
    r2.font.size = Pt(11)
    r3 = p.add_run(detail)
    r3.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("FUSION FLOW V2 — BKD")
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = DARK_BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Standard Operating Procedure")
run2.font.size = Pt(18)
run2.font.color.rgb = MID_BLUE

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Synovia Digital Ltd  ·  Customs & Supply Chain Integration  ·  June 2026")
run3.font.size = Pt(11)
run3.font.color.rgb = GREY
run3.font.italic = True

doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. System Overview")
body("Fusion Flow V2 is Synovia Digital's customs declaration management portal for Birkdale (BKD). It manages the full TSS workflow: ENS headers, consignments, goods items, SFDs, SDIs and GMR/GVMS Route A movements.")

heading2("1.1 Key Workflow")
step("1", "Email received", "Birkdale sends a sales order Excel via email. Fusion fetches it automatically via Microsoft Graph.")
step("2", "ENS created", "The Excel is parsed and staged into an ENS header with consignments and goods.")
step("3", "TSS submission", "ENS is submitted to TSS. TSS processes and issues an SFD or EIDR reference.")
step("4", "SDI discovered", "When the consignment is authorised for movement, TSS generates an SDI. Fusion discovers and stages it.")
step("5", "SDI submitted", "The SDI is submitted to TSS before the monthly deadline. TSS validates and closes it.")

heading2("1.2 Tenants")
add_table(
    ["Code", "Name", "Login", "Environment"],
    [
        ["BKD", "Birkdale", "birkdale / admin", "PRODUCTION"],
        ["SYD", "Synovia Digital", "synovia / admin", "DEMO / QAS"],
        ["CWF", "Countrywide", "countrywide / admin", "—"],
        ["CLR", "Clarity Cargo", "claritycargo / admin", "—"],
    ]
)

heading2("1.3 Databases")
add_table(
    ["Database", "Purpose"],
    [
        ["Fusion_TSS_Automation_PRD", "Live production — BKD tenant"],
        ["Fusion_TSS_Automation_QAS", "QAS/testing — copy of PRD, SYD tenant in demo mode"],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. LOGIN
# ══════════════════════════════════════════════════════════════════════════════
heading1("2. Logging In")
body("Navigate to the Fusion Flow portal URL. Enter your tenant username and password and click Sign In.")
doc.add_paragraph()
screenshot("01_login.png", "Login screen — enter tenant credentials (e.g. birkdale / admin for BKD production)")
note("The top-right badge shows the active TSS environment. Always verify it shows PRODUCTION before submitting real declarations.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. NAVIGATION
# ══════════════════════════════════════════════════════════════════════════════
heading1("3. Navigation & Layout")
body("The top bar is present on every page. It contains the Fusion Flow logo, main navigation links, the TSS environment badge and the tenant brand name.")
doc.add_paragraph()
screenshot("ui/nav_bar.png", "Top navigation bar — logo, nav links, TSS environment badge (PRODUCTION), tenant branding")

heading2("3.1 Navigation Items")
add_table(
    ["Link", "Purpose"],
    [
        ["Dashboard", "Automation activity summary — email counts, TSS sync results, GVMS authorisations, record totals by status"],
        ["ENS", "Entry Summary Declarations — headers submitted to TSS for GVMS Route A movements"],
        ["Consignments", "Individual DEC consignments — linked to ENS, SFD, SDI and goods items"],
        ["SFD", "Simplified Frontier Declarations / EIDR records issued by TSS after consignment clearance"],
        ["SDI", "Supplementary Declarations — post-clearance declarations submitted to TSS monthly"],
        ["Ingestion", "Email inbox log, document uploads and Excel batch staging controls"],
        ["Technical Logs", "Full audit trail: API calls, email intake, notifications, status changes, errors, job runs"],
        ["Master Data", "Company, partners, products, TSS choice values, EORI checker, validation settings"],
        ["Settings", "TSS API credentials, environment toggles, runtime configuration in AppConfiguration"],
        ["Logout", "End the current session and return to the login screen"],
    ]
)

heading2("3.2 TSS Environment Badge")
body("Shown top-right of the nav bar. Colour-coded to immediately flag the active environment:")
add_table(
    ["Badge", "Colour", "Meaning"],
    [
        ["PRODUCTION", "Dark green", "Live TSS API, real declarations — Render IPs required"],
        ["DEMO", "Teal", "Demo mode, no live TSS calls sent from the app"],
        ["TEST", "Orange", "TSS test environment (api.tsstestenv.co.uk)"],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. BUTTON COLOURS
# ══════════════════════════════════════════════════════════════════════════════
heading1("4. Button & Colour Reference")
body("Every button in Fusion Flow follows a consistent colour convention. Learning the colour means knowing the action type at a glance — no need to read the label first.")

add_table(
    ["Style", "Colour", "Used for", "Examples"],
    [
        ["btn-primary", "Blue solid", "Main positive action on a page", "New ENS, Submit ready SDIs, Stage Batch, Continue"],
        ["btn-outline-secondary", "Grey outline", "Secondary or navigate-back actions; also the More dropdown trigger", "Refresh, More, CSV, Back to SFD list, ENS Header"],
        ["btn-info", "Teal solid", "Immediate sync from TSS (updates status right now)", "Sync TSS Now (consignment detail)"],
        ["btn-outline-info", "Teal outline", "Query or discover records from TSS without mutating data", "Discover SDIs"],
        ["btn-danger", "Red solid", "Permanent local delete (irreversible in bulk)", "Delete (ENS Select mode)"],
        ["btn-outline-danger", "Red outline", "Softer destructive; same effect, lower visual weight", "Delete (SDI Select mode)"],
        ["btn-outline-primary", "Blue outline", "Navigation link to a related record (not an action)", "SFD DEC link, Parent DEC on SFD"],
        ["dropdown-item text-success", "Green text in dropdown", "Submit action inside the More menu", "Submit this SDI"],
        ["dropdown-item text-danger", "Red text in dropdown", "Destructive action inside the More menu", "Cancel in TSS"],
    ]
)

heading2("4.1 The 'More' Dropdown")
body("Pages with many actions put secondary operations inside a More button to keep the toolbar uncluttered. Click it to reveal all available actions. Items in green text submit data to TSS; items in red text are destructive or cancel operations.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. STATUS BADGES
# ══════════════════════════════════════════════════════════════════════════════
heading1("5. Status Badge Reference")
body("Status badges appear on list rows and detail pages. Many records show two badges: Local Status (what Fusion staged) and TSS Status (what TSS confirmed). The TSS badge is always authoritative.")

add_table(
    ["Colour", "Meaning", "Statuses using it"],
    [
        ["Green", "Complete or accepted by TSS", "CLOSED, AUTHORISED FOR MOVEMENT, SUBMITTED"],
        ["Teal", "In-progress — TSS has registered but not finalised", "ARRIVED, CREATED (TSS goods ID), N TOTAL count badges"],
        ["Grey", "Pending — not yet submitted or awaiting response", "DRAFT, PENDING REVIEW"],
        ["Orange", "Requires user action in the TSS portal — API cannot resolve", "TRADER INPUT REQUIRED"],
        ["Red", "Error or blocked — urgent action required", "BLOCKED (autosubmit failed), OVERDUE (past deadline)"],
        ["Dark grey", "Terminal — no further action possible", "CANCELLED"],
        ["Purple", "Awaiting payment or duty resolution", "PENDING PAYMENT"],
    ]
)
note("When a row shows two badges (Local + TSS), the TSS badge is authoritative. Local shows what Fusion staged; TSS shows what HMRC/TSS recorded.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
heading1("6. Dashboard")
body("The dashboard shows the current automation activity: emails received, documents staged, TSS API sync results, GVMS authorisations, and summary counts for ENS, consignments, goods and SDIs.")
doc.add_paragraph()
screenshot("02_dashboard.png", "Dashboard — Automation Activity Logs with ING mailbox, documents, TSS sync and GVMS counts")

heading2("Key Indicators")
add_table(
    ["Widget", "What it shows", "Action if wrong"],
    [
        ["ING · MAILBOX", "Emails fetched from Graph mailbox", "Check Graph settings in Admin > Settings"],
        ["ING · DOCUMENTS", "Excel files saved from emails", "Check Ingestion tab for errors"],
        ["TSS · API SYNC", "Consignments synced from TSS", "Check Technical Logs > API Calls"],
        ["GVMS · MOVEMENT AUTH", "ENS authorised for movement", "Check ENS list, verify TSS status"],
        ["SUPPLEMENTARY DECS (SDI)", "Total SDIs by status", "Check SDI list — TRADER INPUT REQUIRED needs action"],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. ENS
# ══════════════════════════════════════════════════════════════════════════════
heading1("4. ENS Declarations")
body("ENS (Entry Summary Declarations) are created from inbound sales order Excels. Each ENS contains consignments and goods items submitted to TSS.")
doc.add_paragraph()
screenshot("03_ens_list.png", "ENS Declarations — list view showing all headers with TSS status, arrival date and cargo counts")

heading2("4.1 ENS List Buttons")
add_table(
    ["Button / Control", "Action"],
    [
        ["Refresh", "Reload the ENS list from the database"],
        ["New ENS", "Open the create form to manually create a new ENS header"],
        ["More > Import from TSS", "Import an existing ENS from a TSS ENS reference"],
        ["More > Bulk CSV Export", "Download all visible ENS records as CSV"],
        ["Select mode", "Toggle checkbox select — enables local cleanup and CSV export for selected rows"],
        ["Status tabs (ALL / DRAFT / ARRIVED / ...)", "Filter the list by TSS or local status"],
        ["Search box", "Filter by ENS ref, carrier, EORI, port or transport identity"],
        ["Row click", "Open the ENS detail page for that header"],
    ]
)

heading2("4.2 ENS Status Flow")
add_table(
    ["Local Status", "TSS Status", "Meaning"],
    [
        ["DRAFT", "—", "Created locally, not yet submitted to TSS"],
        ["SUBMITTED", "ARRIVED", "Goods have arrived at port"],
        ["SUBMITTED", "AUTHORISED FOR MOVEMENT", "GVMS movement authorised — SDI may follow"],
        ["IMPORTED", "—", "Imported from TSS reference, not staged locally"],
    ]
)

heading2("4.3 ENS Detail Page")
screenshot("11_ens_detail.png", "ENS Detail — header fields, status badges, consignments panel and action buttons")
add_table(
    ["Button / Control", "Action"],
    [
        ["More > Edit", "Edit the ENS header fields (carrier, arrival date, transport identity, etc.)"],
        ["More > Submit to TSS", "Submit this ENS to TSS (only valid when in DRAFT status)"],
        ["More > Sync from TSS", "Force a TSS status sync for this ENS header"],
        ["More > Import Consignments", "Import child consignments from a TSS DEC reference into this ENS"],
        ["More > Delete (local)", "Remove the ENS from local staging — does NOT cancel in TSS"],
        ["Consignment links", "Click any DEC reference to open that consignment's detail page"],
        ["SFD links", "Click any SFD/EIDR reference to open the SFD detail page"],
    ]
)

heading2("4.4 Creating a New ENS")
step("1", "Click 'New ENS'", "Top right button, or let automation create it from an inbound email.")
step("2", "Fill header details", "Movement type, port (GBAUBELBELBEL), carrier, arrival date/time.")
step("3", "Add consignments and goods", "Or import from a TSS DEC reference.")
step("4", "Submit to TSS", "Use More > Submit on the ENS detail page.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. CONSIGNMENTS
# ══════════════════════════════════════════════════════════════════════════════
heading1("5. Consignments")
body("Consignments are child records of an ENS header. Each DEC reference in TSS maps to one consignment, with its goods items.")
doc.add_paragraph()
screenshot("04_cons_list.png", "Consignments — STG view with linked ENS, SFD, SDI, goods count and TSS status")

heading2("5.1 Consignment List Buttons")
add_table(
    ["Button / Control", "Action"],
    [
        ["New Consignment", "Create a consignment manually and link it to an ENS header"],
        ["More > Bulk CSV Export", "Export all visible consignments as CSV"],
        ["More > Bulk Delete (local)", "Remove selected consignments from local staging only"],
        ["Select mode", "Toggle checkbox select for bulk operations"],
        ["Status tabs", "Filter by TSS status (ARRIVED, AUTHORISED FOR MOVEMENT, TRADER INPUT REQUIRED, etc.)"],
        ["Search box", "Filter by DEC ref, sales order, EORI or TSS message"],
        ["Row click", "Open the consignment detail page"],
    ]
)

heading2("5.2 Consignment List Key Columns")
add_table(
    ["Column", "Meaning"],
    [
        ["DEC REF", "TSS DEC reference (e.g. DEC000000017101535)"],
        ["SFD", "Simplified Frontier Declaration reference issued by TSS"],
        ["SDI", "Supplementary Declaration reference (if issued)"],
        ["GOODS", "Number of goods items staged locally"],
        ["DOC / REF", "Transport document / sales order reference"],
    ]
)

heading2("5.3 Consignment Detail Page & Goods Items")
screenshot("12_cons_detail.png", "Consignment Detail — header fields, SFD link, action buttons and goods items table with TSS IDs")
add_table(
    ["Button / Control", "Action"],
    [
        ["Sync TSS Now", "Force an immediate TSS status sync for this consignment and its goods"],
        ["More > Edit", "Edit consignment fields (consignor, consignee, doc number, etc.)"],
        ["More > Submit / Update to TSS", "Submit (DRAFT) or update (existing DEC) this consignment to TSS"],
        ["More > Add Good Item", "Add a new goods line to this consignment"],
        ["More > Delete (local)", "Remove the consignment from local staging — does NOT cancel in TSS"],
        ["SFD link (top)", "Navigate to the linked SFD/EIDR detail page"],
        ["Goods row > Edit", "Edit individual goods item fields (taric, weight, package type, invoice amount)"],
        ["Goods row > Delete", "Remove a single goods item from local staging"],
    ]
)
note("Goods items with TSS ID populated have been synced or submitted to TSS. Items with no TSS ID are local-only and need submission.")
note("'No SDI yet' on a consignment that is AUTHORISED FOR MOVEMENT means TSS has not generated the SDI. Run SDI Discovery from the SDI list page.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. SDI
# ══════════════════════════════════════════════════════════════════════════════
heading1("6. Supplementary Declarations (SDI)")
body("SDIs are the post-clearance customs declarations required for EIDR/SFD movements. They must be submitted to TSS before the monthly deadline (10th of the following month).")
doc.add_paragraph()
screenshot("05_sdi_list.png", "Supplementary Declarations — deadline countdown, overdue count and status tabs")

heading2("6.1 SDI List Buttons")
add_table(
    ["Button / Control", "Action"],
    [
        ["Submit ready SDIs", "Batch-submit all DRAFT SDIs that have complete goods data to TSS"],
        ["Discover SDIs", "Query TSS for new SUP references on AUTHORISED consignments — stages new SDI records"],
        ["More > Import SDI", "Manually import an SDI from a known SUP reference"],
        ["More > Bulk CSV Export", "Export visible SDI records as CSV"],
        ["More > Sync all from TSS", "Force TSS status sync for all SDIs in the current filter"],
        ["Status tabs (ALL / DRAFT / CLOSED / TRADER INPUT REQUIRED / CANCELLED)", "Filter list by TSS status"],
        ["SDI deadline banner", "Shows days remaining to the monthly cutoff (10th of following month)"],
        ["Search box", "Filter by SUP ref, SFD, DEC, EORI or TSS message text"],
        ["Row click", "Open SDI detail page"],
    ]
)

heading2("6.2 SDI Detail Page")
screenshot("14_sdi_detail.png", "SDI Detail — declaration chain, TSS exchange log, BLOCKED/PENDING REVIEW state and More actions")
add_table(
    ["Button / Control", "Action"],
    [
        ["More > Submit to TSS", "Submit or re-submit this SDI to TSS (all goods items must have invoice amounts)"],
        ["More > Sync from TSS", "Force a status sync — fetches latest TSS response for this SUP reference"],
        ["More > Edit SDI Header", "Edit header-level fields (importer EORI, exporter EORI, procedure codes)"],
        ["More > Edit Goods Items", "Navigate to goods items edit — update invoice amounts, weights, taric codes"],
        ["More > Reject (local)", "Mark SDI as locally rejected — does NOT cancel in TSS"],
        ["TSS response > Details", "Expand the TSS JSON exchange — shows the exact error message from TSS"],
        ["Declaration chain links", "Jump to the linked SFD or DEC consignment detail page"],
    ]
)
note("BLOCKED means autosubmit cannot proceed. TRADER INPUT REQUIRED means the TSS portal itself needs manual action — the API cannot resolve it.", danger=True)

heading2("6.3 SDI Status Reference")
add_table(
    ["Status", "Meaning", "Action"],
    [
        ["DRAFT", "Staged locally, not submitted to TSS", "Review goods have invoice amounts, then submit"],
        ["PENDING REVIEW", "TSS returned validation errors", "Check TSS message column, fix values, re-submit"],
        ["TRADER INPUT REQUIRED", "TSS requires manual action in TSS portal", "Log into TSS portal to resolve — API cannot fix TIR"],
        ["CLOSED", "TSS accepted and finalised the declaration", "No action needed"],
        ["CANCELLED", "Cancelled in TSS or locally", "No action needed"],
    ]
)

heading2("6.4 Submitting SDIs")
step("1", "Click 'Discover SDIs'", "Find new SUP references from TSS for arrived consignments.")
step("2", "Review DRAFT SDIs", "Check each goods item has a valid invoice amount and commodity code.")
step("3", "Click 'Submit ready SDIs'", "Submits all goods-complete SDIs to TSS in one batch.")
step("4", "Monitor status", "Cron sync refreshes TSS status every hour. Manual sync on each SDI detail page.")

note("Monthly Deadline: SDIs for consignments that arrived in the previous calendar month must be submitted by the 10th. The SDI list shows a countdown and highlights overdue entries.", danger=True)

heading2("6.5 1UKI Document (Portland Cement)")
body("All goods items containing commodity code 2523290000 (Portland cement) require a 1UKI document.")
body("Reference: XIUKIM37969209200020230711113853 — Birkdale's fixed UKIM authorisation number.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. SFD
# ══════════════════════════════════════════════════════════════════════════════
heading1("7. Simplified Frontier Declarations (SFD)")
body("SFDs are issued by TSS automatically when a consignment clears. Most Birkdale consignments use EIDR rather than a full SFD.")
doc.add_paragraph()
screenshot("06_sfd_list.png", "SFD/EIDR list — path type, ENS consignment link and TSS authorisation status")
add_table(
    ["Path", "Meaning"],
    [
        ["EIDR", "Electronic Import Duty Relief — most common for Birkdale"],
        ["SFD", "Simplified Frontier Declaration — issued when TSS creates a formal SFD reference"],
    ]
)

heading2("7.1 SFD List Buttons")
add_table(
    ["Button / Control", "Action"],
    [
        ["Refresh", "Reload the SFD list"],
        ["More > Sync from TSS", "Force a TSS sync for all visible SFDs"],
        ["Status tabs", "Filter by TSS status (ARRIVED, AUTHORISED FOR MOVEMENT, etc.)"],
        ["Search box", "Filter by DEC ref, MRN/EIDR, EORI, sales order"],
        ["Row click", "Open SFD detail page"],
    ]
)

heading2("7.2 SFD Detail Page")
screenshot("13_sfd_detail.png", "SFD Detail — consignment fields, MRN/EIDR, links to parent DEC and ENS header")
add_table(
    ["Button / Control", "Action"],
    [
        ["Parent DEC", "Navigate to the parent consignment (DEC reference) that owns this SFD"],
        ["ENS Header", "Navigate to the ENS header that contains the parent consignment"],
        ["Back to SFD list", "Return to the SFD list view"],
    ]
)
body("SFD detail is read-only. All data is sourced from TSS via sync — edit the parent consignment to make changes.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. INGESTION
# ══════════════════════════════════════════════════════════════════════════════
heading1("8. Email Automation & Ingestion")
body("Fusion automatically fetches emails from the Birkdale mailbox via Microsoft Graph, extracts Excel sales order attachments and stages ENS records.")
doc.add_paragraph()
screenshot("07_ingestion.png", "Email Automation — inbox log showing received emails, attachments and staging status")

heading2("8.1 How It Works")
step("1", "Email arrives", "Birkdale sender emails nexus@synoviaflow.cloud with a sales order Excel attached.")
step("2", "Fetch New Emails", "Fusion polls Graph API and logs each email. Attachments are saved to the files/ folder with received date in filename.")
step("3", "Parser runs", "Excel is parsed into ING.BKD_SalesOrderLine records (one per goods item).")
step("4", "Stage Batch", "ENS header, consignments and goods are created from the parsed data.")

heading2("8.2 Manual Upload")
body("Files can also be drag-dropped into the upload area and staged manually. Use Template (no SD) or Template (with SD) to download the correct Excel template.")
note("Creation mode: 'Auto-create clean' creates clean ENS records. Change to 'Review required' to stage into draft before submission.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. MASTER DATA
# ══════════════════════════════════════════════════════════════════════════════
heading1("9. Master Data")
body("Reference data used across all declarations. Changes here affect validation and auto-population of party details.")
doc.add_paragraph()
screenshot("08_master_data.png", "Master Data — company, partners, products, TSS choice values and EORI validation")
add_table(
    ["Section", "Contents", "Current count"],
    [
        ["Company", "Tenant company details and ingest defaults", "Configured"],
        ["Partners", "Importers, exporters, consignors, consignees, carriers", "221 active"],
        ["Products", "Commodity catalogue — taric codes, weights, package defaults", "7,932 products"],
        ["TSS Choice Values", "CV_* reference tables synced from TSS", "58 tables"],
        ["EORI Checker", "Bulk GB/XI/EU EORI validation", "Up to 100 at once"],
    ]
)
note("Strict Masterdata Validation is currently ON for BKD. Disable temporarily for end-to-end tests with demo values.", danger=False)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. SETTINGS
# ══════════════════════════════════════════════════════════════════════════════
heading1("10. Admin Settings")
body("Manage TSS API credentials, email configuration and runtime toggles. All values are stored in BKD.AppConfiguration and editable without a deployment.")
doc.add_paragraph()
screenshot("09_settings.png", "Admin Settings — overview of all 7 configuration sections")
note("Click 'Save All Settings' after any change. Individual sections have a 'Test Connection' or 'Test SMTP' button to verify credentials immediately.")
doc.add_paragraph()

heading2("10.1 TSS Portal API")
body("Credentials and endpoints for the Trader Support Service REST API. All live submissions, syncs and validations route through these settings.")
screenshot("settings/tss_api.png", "TSS Portal API — environment, credentials and endpoint URLs")
add_table(
    ["Key", "What it does"],
    [
        ["BASE_URL", "Production TSS endpoint. Only reachable from Render deployment IPs."],
        ["TEST_URL", "QAS/test endpoint (api.tsstestenv.co.uk). Switch ENVIRONMENT to 'test' to use this."],
        ["ENVIRONMENT", "CRITICAL: production / test / demo. Demo simulates responses, makes no real TSS calls. Always verify before going live."],
        ["USER", "TSS API username (e.g. API.TSS0012045). Obtained from TSS account."],
        ["PASSWORD", "TSS API password (masked). Contact Synovia to rotate."],
        ["ACT_AS", "Optional TSS customer_account_sys_id for delegated/bureau API calls. Leave blank if not needed."],
    ]
)
note("The nav bar shows the active environment badge. Verify PRODUCTION is shown before submitting live declarations. Use 'demo' for local testing — no Render IPs required.", danger=True)
doc.add_paragraph()

heading2("10.2 Email / SMTP")
body("Outbound email for SDI deadline alerts and all automation notification emails.")
screenshot("settings/email_smtp.png", "Email / SMTP — outbound email credentials")
add_table(
    ["Key", "What it does"],
    [
        ["SERVER", "SMTP server hostname (e.g. smtp.office365.com or smtp.gmail.com)."],
        ["PORT", "SMTP port — typically 587 (STARTTLS) or 465 (SSL)."],
        ["SENDER_EMAIL", "From address used for all system-generated emails."],
        ["SENDER_PASSWORD", "SMTP authentication password (masked)."],
        ["ALERT_RECIPIENT", "Default recipient for deadline alerts and staging failure notifications."],
    ]
)
note("Use the 'Test SMTP' button to send a test email and verify credentials after any change.")
doc.add_paragraph()

heading2("10.3 Inbound Email / Microsoft Graph")
body("Configures the mailbox that Fusion Flow polls for inbound supplier emails. Uses an Azure AD App Registration — no user login required.")
screenshot("settings/graph_mail.png", "Inbound Email / Microsoft Graph — mailbox and polling configuration")
add_table(
    ["Key", "What it does"],
    [
        ["ENABLED", "Master ON/OFF for Graph mail polling. Set false to stop all email ingestion immediately. Can be overridden by a Render environment variable."],
        ["TENANT_ID", "Azure AD tenant ID of the Microsoft 365 organisation that owns the mailbox."],
        ["CLIENT_ID", "App registration client ID. Needs Mail.Read + Mail.ReadWrite Graph permissions with admin consent."],
        ["CLIENT_SECRET", "App registration secret (masked). Rotate in Azure portal; update here immediately after."],
        ["MAILBOX", "Email address of the shared mailbox to poll (e.g. operations@birkdale.com)."],
        ["FOLDER", "Mailbox folder to scan — INBOX or a named subfolder (e.g. Inbox/BKD)."],
        ["UNREAD_ONLY", "If true, only unread messages are processed. Prevents re-processing already-handled emails."],
        ["MAX_MESSAGES", "Maximum messages per poll cycle. Default 50. Lower if batches are slow."],
        ["PROCESSED_FOLDER", "After processing, messages are moved here (e.g. Processed). Blank = mark read only."],
    ]
)
note("Use the 'Test Graph' button to connect with current credentials and confirm how many messages are waiting.")
doc.add_paragraph()

heading2("10.4 Invoice Auto-Staging")
body("Controls how inbound emails are converted into ENS records. Default values fill in mandatory TSS fields that cannot be derived from the supplier invoice alone.")
screenshot("settings/auto_staging.png", "Invoice Auto-Staging — mode, defaults and transport/carrier fields")
add_table(
    ["Key", "Default", "What it does"],
    [
        ["ENABLED", "true", "Master toggle. When false, emails are logged but no staging records are created."],
        ["MODE", "review_required", "review_required = operator approves on Ingestion screen before ENS is created. auto_create_if_clean = ENS created immediately when validation passes with zero errors."],
        ["AUTO_VALIDATE", "true", "Run pipeline validation immediately after staging. Errors appear inline on the review screen."],
        ["SUPPLIER_NAME", "(blank)", "Expected supplier name. Used to filter emails. Leave blank to accept all senders."],
        ["DEFAULT_MOVEMENT_TYPE", "1", "Mode of transport: 1=Sea, 3=Road, 4=Air."],
        ["DEFAULT_ARRIVAL_PORT", "(blank)", "UN/LOCODE arrival port (e.g. GBHUL). Required for TSS submission."],
        ["DEFAULT_ARRIVAL_HOURS_AHEAD", "4", "When no arrival date in email, arrival = now + N hours."],
        ["DEFAULT_CARRIER_EORI", "(blank)", "Carrier EORI — filled from email if present, falls back to this default."],
        ["DEFAULT_IMPORTER_EORI", "(blank)", "Falls back to CompanyMaster EORI when blank."],
        ["DEFAULT_PROCEDURE_CODE", "4000", "Customs procedure code for all goods items."],
        ["DEFAULT_INVOICE_CURRENCY", "GBP", "Invoice currency for goods valuation."],
        ["DEFAULT_COUNTRY_OF_ORIGIN", "GB", "Country of origin applied to all staged goods items."],
        ["DEFAULT_PACKAGE_TYPE", "PK", "Package type code (PK = packages)."],
        ["DEFAULT_VALUATION_METHOD", "1", "Customs valuation method — 1 = transaction value."],
    ]
)
note("Carrier, transport and importer fields default to blank and fall back to CompanyMaster or email-parsed values. Fill in only the fields that are constant for every shipment.")
doc.add_paragraph()

heading2("10.5 SDI / SupDec Automation")
body("Controls the one-step Supplementary Declaration worker. After TSS exposes the SFD/SUP relationship, the system creates and optionally submits the SupDec automatically.")
screenshot("settings/sdi_automation.png", "SDI / SupDec Automation — defaults and AUTOSUBMIT kill switch")
add_table(
    ["Key", "Default", "What it does"],
    [
        ["AUTOSUBMIT TSS", "KILL SWITCH", "When blank/false, the SDI worker stages records locally but does NOT submit to TSS. Enable only after SDI defaults are fully configured and tested."],
        ["MAX_ITEMS", "integer", "Max goods items per SupDec before system warns or splits."],
        ["DEFAULT_REPRESENTATION_TYPE", "3", "TSS representation type: 3 = direct representation (importer = declarant)."],
        ["DEFAULT_INCOTERM", "DDP", "Default Incoterms code for all SupDec goods (e.g. DDP, CIF, FOB)."],
        ["DEFAULT_POSTPONED_VAT", "no", "Whether to apply Postponed VAT Accounting by default. yes/no."],
        ["DEFAULT_GOODS_DOMESTIC_STATUS", "D", "Domestic/non-domestic status code for SDI goods items."],
        ["DEFAULT_MOVEMENT_TYPE", "3", "Movement type for SDI records (3 = RoRo / standard import)."],
        ["DEFAULT_NATURE_OF_TRANSACTION", "11", "Nature of transaction code for SupDec headers."],
        ["DEFAULT_NI_ADDITIONAL_INFORMATION_CODES", "NIREM", "NI additional information codes. NIREM = NI goods not at risk."],
    ]
)
note("AUTOSUBMIT TSS is the primary SDI kill switch. Disable immediately if there is any concern about data quality or TSS connectivity. Existing in-flight records are not affected.", danger=True)
doc.add_paragraph()

heading2("10.6 Validation Controls")
body("Runtime switches that control how strictly local pipeline validation blocks before sending data to TSS.")
screenshot("settings/validation.png", "Validation Controls — strict master-data validation toggle")
add_table(
    ["Key", "Default", "What it does"],
    [
        ["STRICT_MASTERDATA_VALIDATION", "true", "When true, blocks submission if partners or products are not in Master Data. Set false only for end-to-end testing with demo/random values. Must be true in production."],
    ]
)
note("Never leave strict validation disabled in production. Unrecognised EORIs and unmapped products will pass through to TSS unchecked.", danger=True)
doc.add_paragraph()

heading2("10.7 Email Automation Notifications")
body("Controls what automated notification emails are sent and to whom at each workflow event.")
screenshot("settings/notifications.png", "Email Automation Notifications — event-based email settings")
add_table(
    ["Setting", "When it fires", "Who receives it"],
    [
        ["ENS_RECEIVED_ENABLED", "New ENS/consignments staged from an inbound email", "ALERT_RECIPIENT"],
        ["CONSIGNMENTS_RECEIVED_ENABLED", "New consignment batch processed from an email attachment", "ALERT_RECIPIENT"],
        ["STAGING_FAILURES_ENABLED", "Staging failed for one or more records", "STAGING_FAILURES_TO (or ALERT_RECIPIENT if blank)"],
        ["MOVEMENT_AUTHORISED_ENABLED", "GVMS authorises a movement — goods cleared to move", "MOVEMENT_AUTHORISED_TO + CC to MOVEMENT_AUTHORISED_CC"],
        ["EMAIL_AUTOMATION_TEST_TO", "Smoke-test triggered by 'Test' button", "Address entered here — use to verify SMTP end-to-end"],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. TECHNICAL LOGS
# ══════════════════════════════════════════════════════════════════════════════
heading1("11. Technical Logs")
body("Full audit trail of email intake, API calls, notifications, status changes, errors and job runs.")
doc.add_paragraph()
screenshot("10_technical.png", "Technical Logs — 2388 API OK calls, 15 failures, 159ms average latency in 24h")
add_table(
    ["Tab", "Contents"],
    [
        ["Email Ingestion", "Per-email staging trace — document, source, status, staged records"],
        ["API Calls", "Every TSS API exchange — call type, HTTP status, payload, response, duration"],
        ["Notifications", "Automated email notifications sent to operators"],
        ["Status Changes", "TSS status transitions logged per entity"],
        ["Local Errors", "Application errors captured at runtime"],
        ["Latency", "API response time trends"],
        ["Job Runs", "Cron job execution history"],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. ENVIRONMENTS
# ══════════════════════════════════════════════════════════════════════════════
heading1("12. Environments (PRD / QAS / SYD)")
add_table(
    ["Environment", "Database", "TSS", "Use for"],
    [
        ["PRODUCTION — BKD", "Fusion_TSS_Automation_PRD", "api.tradersupportservice.co.uk (Render IPs only)", "Live Birkdale operations"],
        ["QAS — SYD", "Fusion_TSS_Automation_QAS", "demo mode or api.tsstestenv.co.uk", "Testing, onboarding, demos"],
    ]
)

heading2("12.1 Switching to QAS locally")
step("1", "Change DB_CONN_STR in .env", "Replace Fusion_TSS_Automation_PRD with Fusion_TSS_Automation_QAS.")
step("2", "Start the app", r".\run_local.ps1 -Service main")
step("3", "Log in as synovia / admin", "SYD tenant, demo mode active, no live TSS calls.")

heading2("12.2 SYD Tenant Database Objects")
add_table(
    ["Schema", "Tables", "Purpose"],
    [
        ["SYD.*", "AppConfiguration, CompanyMaster, Partners, Products", "Per-tenant config and master data"],
        ["STG.SYD_*", "SDI_Headers, GoodsItems, ENS_Headers, ENS_Consignments, SFD_Tracking, GMR_Movements", "Pipeline staging tables"],
        ["ING.SYD_*", "SalesOrderLine, EmailMessage, SourceFileLog, EmailAttachment", "Ingestion source tables"],
        ["TSS.SYD_*", "API_Exchanges, SDI_Headers, SDI_GoodsItems, ENS_Headers, GoodsItems", "TSS mirror tables"],
    ]
)
note("SQL is automatically rewritten per tenant via tenantize_sql() — STG.BKD_Table becomes STG.SYD_Table when logged in as synovia. No code changes needed per route.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13. AUTOMATION & SETTINGS GUIDE
# ══════════════════════════════════════════════════════════════════════════════
heading1("13. Automation & Settings Guide")
body("Fusion Flow automates the full customs chain from inbound supplier email through to ENS submission and Supplementary Declaration creation. Human review is only required where data is incomplete or a declaration fails TSS validation.")

heading2("13.1 End-to-End Automation Flow")
step("1", "Inbound email arrives", "Supplier sends a Sales Order Excel (.xlsx) or PDF invoice to the configured mailbox.")
step("2", "Microsoft Graph poll", "A scheduled cron job polls the mailbox via the Graph API (settings: Inbound Email / Microsoft Graph). Up to MAX_MESSAGES unread messages are fetched each cycle.")
step("3", "Email classified", "Each message is classified by its content: Type A = carrier DETAILS body (creates ENS header), Type B = Sales Orders XLSX attachment (adds consignments + goods), Type C = PDF/CSV invoice batch (one consignment per attachment), Type D = no relevant content (skipped).")
step("4", "ENS staged", "Defaults from Invoice Auto-Staging settings fill in mandatory TSS fields. Mode controls whether the operator must review before an ENS is created (review_required) or whether it is created immediately (auto_create_if_clean).")
step("5", "Submit pipeline", "ENS header, consignments and goods are submitted to TSS. TSS issues DEC references back.")
step("6", "Sync pipeline", "A cron sync reads TSS status back into local staging tables. Notification emails fire when configured events occur.")
step("7", "SDI discovery", "Once a consignment is AUTHORISED FOR MOVEMENT, the SDI worker queries TSS for the linked SUP reference and creates local SDI staging records.")
step("8", "SDI auto-submit", "If the AUTOSUBMIT TSS kill switch is enabled, SupDecs are sent to TSS automatically. Failures are logged and the record moves to TRADER INPUT REQUIRED.")
doc.add_paragraph()

heading2("13.2 Email Types Explained")
add_table(
    ["Type", "What it contains", "What Fusion does"],
    [
        ["A — DETAILS (body only)", "Carrier block in email body: vessel, arrival port, arrival date/time", "Creates or updates the ENS header draft"],
        ["B — Sales Orders XLSX", ".xlsx attachment following BKD Sales Orders format", "Creates consignments + goods items linked to the ENS header"],
        ["C — Invoice batch", ".pdf / .csv / .zip attachments", "One ENS + one consignment per attachment, goods from invoice lines"],
        ["D — Skip", "No supported attachments, no carrier block", "Message marked read / moved, no staging action taken"],
    ]
)
note("Sender filtering: ALLOWED_SENDER_DOMAINS in Graph settings limits which senders are processed. Default for BKD: birkdalesales.com. Emails from other senders are silently skipped.")

heading2("13.3 Kill Switches — Quick Reference")
add_table(
    ["To stop", "Setting to change", "Set to"],
    [
        ["All email polling", "Inbound Email / Graph → ENABLED", "false"],
        ["ENS staging from emails", "Invoice Auto-Staging → ENABLED", "false"],
        ["Hold ENS for review", "Invoice Auto-Staging → MODE", "review_required"],
        ["SDI auto-submission to TSS", "SDI / SupDec Automation → AUTOSUBMIT TSS", "blank / false"],
        ["Route TSS calls to test env", "TSS Portal API → ENVIRONMENT", "test"],
        ["Fully simulate TSS (no API calls)", "TSS Portal API → ENVIRONMENT", "demo"],
    ]
)
note("After any emergency stop: 1) Resolve the issue. 2) Verify ENVIRONMENT is correct. 3) Confirm STRICT_MASTERDATA_VALIDATION is true. 4) Run a small manual test batch. 5) Re-enable. 6) Monitor Ingestion Queue and Technical Logs.", danger=True)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 14. DAILY CHECKLIST
# ══════════════════════════════════════════════════════════════════════════════
heading1("14. Daily Operations Checklist")

heading2("Every Morning")
for item in [
    "Check Dashboard — verify ING Mailbox count increased overnight.",
    "Check SDI list — any new TRADER INPUT REQUIRED? If yes, log into TSS portal to resolve.",
    "Check SDI deadline countdown — if < 5 days, prioritise any DRAFT or PENDING_REVIEW SDIs.",
    "Check Technical Logs → API Calls — any red failures in last 24h?",
]:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(11)

heading2("When a New Email Arrives")
for item in [
    "Ingestion tab → 'Fetch New Emails' (or wait for cron to run).",
    "Verify the email appears in Email Intake Log with status PROCESSED.",
    "Verify ENS was created in ENS list.",
    "Submit ENS if still in DRAFT status.",
]:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(11)

heading2("Monthly (by the 10th)")
for item in [
    "SDI list → click 'Submit ready SDIs' for all remaining DRAFT SDIs from previous month.",
    "Resolve any TRADER INPUT REQUIRED via TSS portal.",
    "Confirm all previous month SDIs reach CLOSED status.",
]:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(11)

heading2("Key Contacts & Links")
add_table(
    ["Resource", "Details"],
    [
        ["TSS Portal", "https://www.tradersupportservice.co.uk"],
        ["Render Dashboard", "render.com — fusion-flow-bkd services and cron jobs"],
        ["Synovia Support", "alvaro.molina@synoviadigital.com"],
    ]
)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"SOP saved: {OUT}")
